"""
Generate a JCP (Journal of Chemical Physics) Word submission template.
Based on AIP Publishing official Author Instructions (2026).
Includes comprehensive figure AND table requirements.
"""

import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ── Style helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)
pf.space_before = Pt(0)

def add_style(name, font_size, bold=True, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=12, space_after=6, color=None):
    try:
        s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        s = doc.styles[name]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(font_size)
    s.font.bold = bold
    s.font.italic = italic
    if color:
        s.font.color.rgb = RGBColor(*color)
    s.paragraph_format.alignment = align
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after = Pt(space_after)
    s.paragraph_format.line_spacing = 2.0
    return s

add_style('JCP_Title', 16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
add_style('JCP_Author', 12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=0)
add_style('JCP_Affil', 10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
add_style('JCP_AbstractLabel', 12, bold=True, space_before=18, space_after=6)
add_style('JCP_H1', 12, bold=True, space_before=14, space_after=6)
add_style('JCP_H2', 12, bold=True, space_before=10, space_after=4)
add_style('JCP_Note', 10, bold=False, italic=True, color=(100, 100, 100), space_before=2, space_after=2)
add_style('JCP_FigCaption', 11, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
add_style('JCP_TableCaption', 11, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
add_style('JCP_Ref', 12, bold=False, space_before=0, space_after=0)

def note(text):
    p = doc.add_paragraph(style='JCP_Note')
    p.add_run(text)

def body(text):
    return doc.add_paragraph(text, style='Normal')

def h1(text):
    doc.add_paragraph(text, style='JCP_H1')

def h2(text):
    doc.add_paragraph(text, style='JCP_H2')

def fig_placeholder(fig_num, width_note):
    doc.add_paragraph(style='Normal')
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[Insert Figure {fig_num} here]')
    run.italic = True
    run.font.size = Pt(11)
    note(width_note)

def make_table(caption_label, caption_text, headers, rows, footnote_text=None):
    """Create a formatted table with caption, headers, data, and optional footnotes."""
    p = doc.add_paragraph(style='JCP_TableCaption')
    r = p.add_run(caption_label)
    r.bold = True
    p.add_run(caption_text)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    if footnote_text:
        note(footnote_text)
    doc.add_paragraph(style='Normal')

# ══════════════════════════════════════════════════════════════════
# TEMPLATE CONTENT
# ══════════════════════════════════════════════════════════════════

# ── Title ──
doc.add_paragraph(
    'Title of Your Article: A Descriptive Title for the Journal of Chemical Physics',
    style='JCP_Title')

# ── Authors ──
p = doc.add_paragraph(style='JCP_Author')
p.add_run('First Author,')
r = p.add_run('a,*')
r.font.superscript = True
p.add_run(' Second Author,')
r = p.add_run('b')
r.font.superscript = True
p.add_run(' and Third Author')
r = p.add_run('a,c')
r.font.superscript = True

# ── Affiliations ──
for label, text in [
    ('a', 'Department of Chemistry, University Name, City, State Zip Code, Country'),
    ('b', 'Department of Physics, University Name, City, State Zip Code, Country'),
    ('c', 'Corresponding author. Electronic mail: email@university.edu'),
]:
    p = doc.add_paragraph(style='JCP_Affil')
    r = p.add_run(label)
    r.font.superscript = True
    p.add_run(' ' + text)

p = doc.add_paragraph(style='JCP_Affil')
p.paragraph_format.space_before = Pt(6)
p.add_run('(Dated: Month Day, Year)').italic = True

doc.add_paragraph(style='Normal')

# ── Abstract ──
doc.add_paragraph('Abstract', style='JCP_AbstractLabel')
note('[Abstract: single paragraph, ≤250 words. NO equations, footnotes, references, graphics, or tables.]')
body(
    'The abstract should be a single unstructured paragraph of approximately 150–250 words. '
    'It must not contain displayed equations, footnotes, references, graphics, or tables. '
    'Provide a concise summary of the new information, results of general interest, and conclusions.'
)
doc.add_paragraph(style='Normal')

# ══════════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ══════════════════════════════════════════════════════════════════
h1('I. INTRODUCTION')
body(
    'Begin your introduction here. Provide context, explain motivation and significance, '
    'and briefly review relevant prior studies. Use superscript numbers for references.'
)
note('[Manuscript order: Title → Authors → Affiliations → Abstract → Text → Conclusion → '
     'Supplementary Material → Acknowledgments → Author Declarations → Data Availability → '
     'Appendixes → References. All pages numbered consecutively.]')

# ══════════════════════════════════════════════════════════════════
# II. COMPUTATIONAL METHODS
# ══════════════════════════════════════════════════════════════════
h1('II. COMPUTATIONAL METHODS')
h2('A. Subsection Title')
body(
    'Describe computational methods here. Equations must use Word Equation Editor or MathType '
    '(must be editable, NOT images).'
)
note('[Equation rules: punctuated, aligned, numbered on right. Multi-line: operator at beginning of next line. '
     'Use "×" not center dot (except scalar products). Use / for fractions. Math font consistent throughout. '
     'For complex exponents use "exp".]')

p = body('')
p.add_run('Example: ').italic = True
p.add_run('E = mc²').bold = True

h2('B. Another Subsection')
body('Continue describing your methods here...')

# ══════════════════════════════════════════════════════════════════
# III. RESULTS AND DISCUSSION
# ══════════════════════════════════════════════════════════════════
h1('III. RESULTS AND DISCUSSION')
body('Present results and discussion. Reference figures and tables in text.')

# ─────────────────────────────────────────────
# FIGURE REQUIREMENTS (official AIP)
# ─────────────────────────────────────────────
note('')
note('═' * 60)
note('FIGURE REQUIREMENTS (from AIP Author Instructions):')
note('═' * 60)
note('GENERAL RULES:')
note('  - All figures embedded in the text near the callout')
note('  - Do NOT upload separate figure files at initial submission')
note('  - Number in order of appearance: 1, 2, 3, ...')
note('  - Identify parts with (a), (b), etc.')
note('  - All figures must have a caption')
note('  - All figures must include alt text (required)')
note('  - All fonts must be embedded in figures')
note('  - Lettering and labels: avoid large size differences')
note('SIZE LIMITS:')
note('  - One-column max width:  3.37 inches (8.5 cm)')
note('  - Two-column max width:  6.69 inches (17 cm)')
note('  - Max depth:             8.25 inches (21.1 cm)')
note('  - Min label font size:   8-point (2.8 mm high)')
note('  - Min line width:        0.5-point')
note('  - Prepare for 100% reproduction (no large reductions)')
note('RESOLUTION:')
note('  - Line art:              600 dpi, black/white bitmap (NOT grayscale)')
note('  - Halftones (photos):    264 dpi, grayscale (NOT b/w bitmap)')
note('  - Combinations:          600 dpi, grayscale (NOT b/w bitmap)')
note('  - Color online:          300 dpi, TIFF/PS/EPS, RGB mode')
note('FORMATS (priority):')
note('  1. EPS  — best, inline in submission system')
note('  2. PS   — production ready')
note('  3. TIFF — production ready, best for photos')
note('  4. PDF  — only when EPS/TIFF cannot be generated')
note('  5. JPG  — acceptable but NOT inline in system PDF')
note('PRODUCTION PDF GRAPHICS:')
note('  - Shaded/photographic: 600 PPI')
note('  - Line art (no shading): 1200 PPI')
note('  - All fonts embedded; "High Quality Print"')
note('COLOR PRINTING (free online, B&W in print):')
note('  - Color files: .eps, .ps, .tif, or .pdf only')
note('  - ONE version per figure')
note('  - Colors must reproduce well in B&W')
note('  - Text descriptions clear for both print and online')
note('HIGHLIGHT IMAGE (recommended with revised submission):')
note('  - Figure from paper or separate image')
note('  - If not provided, Fig. 1 used automatically')
note('  - Types: EPS, TIFF, JPEG')
note('═' * 60)
note('')

# ── Figure placeholders ──
fig_placeholder(1,
    '[Single-column: max 3.37 in (8.5 cm). EPS/TIFF/PDF. '
    '600 dpi line art / 300 dpi color / 264 dpi halftone. '
    'Fonts embedded. Labels ≥8pt. Lines ≥0.5pt.]')
p = doc.add_paragraph(style='JCP_FigCaption')
r = p.add_run('FIG. 1. '); r.bold = True
p.add_run('Caption. Arabic numerals (1, 2, 3). Parts: (a), (b). '
           'Avoid large size differences between labels.')

fig_placeholder(2,
    '[Two-column: max 6.69 in (17 cm). Wide figures.]')
p = doc.add_paragraph(style='JCP_FigCaption')
r = p.add_run('FIG. 2. '); r.bold = True
p.add_run('Two-column figure. Max width: 6.69 inches (17 cm).')

fig_placeholder(3,
    '[Multi-panel: combine (a),(b) into ONE file. '
    'Label inside figure.]')
p = doc.add_paragraph(style='JCP_FigCaption')
r = p.add_run('FIG. 3. '); r.bold = True
p.add_run('(a) Panel (a). (b) Panel (b). Combine into single file.')

# ─────────────────────────────────────────────
# TABLE REQUIREMENTS (official AIP) — EXPANDED
# ─────────────────────────────────────────────
note('')
note('═' * 60)
note('TABLE REQUIREMENTS (from AIP Author Instructions):')
note('═' * 60)
note('OFFICIAL RULES:')
note('  - Tables should be embedded in the text.')
note('  - Tables should have a caption and be cited in text')
note('    and numbered consecutively, i.e., I, II, III, etc.')
note('  - Footnotes in tables should be labeled as a), b), c), etc.')
note('  - Units should be noted in column headings.')
note('  - If using bold font to emphasize table data, include')
note('    an explanation, i.e., "Boldface denotes..."')
note('  - Unaltered computer output cannot be accepted.')
note('  - All tables must include alt text.')
note('')
note('DETAILED REQUIREMENTS:')
note('  - Numbering:     Roman numerals (I, II, III, ...)')
note('  - Caption:       ABOVE the table, MUST have')
note('  - Position:      Embedded in text')
note('  - Citation:      Must cite in text, numbered consecutively')
note('  - Footnotes:     a), b), c), etc. (NOT numbers!)')
note('  - Units:         In COLUMN HEADINGS (e.g. "Energy (Ha)")')
note('  - Bold font:     MUST explain: "Boldface denotes..."')
note('  - Computer out:  Unaltered output NOT acceptable')
note('  - Alt Text:      REQUIRED for every table')
note('  - Word users:    Use "Insert Table" function')
note('  - Appendix:      Numbering CONTINUES from main text')
note('')
note('TABLE ALT TEXT GUIDELINES:')
note('  - Summarize significant trends')
note('  - Identify main point in relation to your work')
note('  - Do NOT write "table of" (unless specific table type)')
note('  - Do NOT repeat point-by-point data or headings')
note('  - ~25-50 words, 1-2 sentences')
note('  - Be concise, specific, objective')
note('═' * 60)
note('')

# ── TABLE I: Basic table ──
make_table(
    'TABLE I. ',
    'Computed total energies for small molecules at the CCSD(T) level. '
    'Units are noted in column headings.',
    ['System', 'Method', 'Energy (Ha)', 'Error (kcal/mol)'],
    [
        ['H\u2082O', 'CCSD(T)/cc-pVTZ', '-76.3380', '0.00'],
        ['NH\u2083', 'CCSD(T)/cc-pVTZ', '-56.5630', '0.00'],
        ['CH\u2084', 'CCSD(T)/cc-pVTZ', '-40.5180', '0.00'],
    ]
)

# ── TABLE II: Table with footnotes a), b) ──
make_table(
    'TABLE II. ',
    'Comparison of basis set effects on computed properties of H\u2082O. '
    'Units noted in column headings. Footnotes use a), b), c) labels.',
    ['Basis set', 'E (Ha)', 'r_e (\u00c5)', '\u03c9_e (cm\u207b\u00b9)', '\u03bc (D)'],
    [
        ['cc-pVDZ', '-76.2315', '0.962', '3832', '1.86'],
        ['cc-pVTZ', '-76.3380', '0.958', '3854', '1.85'],
        ['cc-pVQZ', '-76.3651', '0.957', '3861', '1.85'],
    ],
    'a) Total electronic energy at equilibrium geometry.\nb) Relative to experimental value.'
)

# ── TABLE III: Table with bold font ──
make_table(
    'TABLE III. ',
    'Benchmark results for different quantum chemistry methods. '
    'Boldface denotes the most accurate result for each property.',
    ['Method', '\u0394E (kcal/mol)', 'MAE (kcal/mol)', 'Max err (kcal/mol)'],
    [
        ['HF', '12.5', '8.3', '15.2'],
        ['B3LYP', '3.2', '2.1', '5.4'],
        ['MP2', '1.8', '1.2', '3.1'],
        ['CCSD(T)', '0.1', '0.1', '0.3'],
    ],
    'Boldface values indicate the lowest error for each column.'
)

# ── TABLE IV: Multi-row table with detailed footnotes ──
make_table(
    'TABLE IV. ',
    'Extended benchmark with multiple systems and footnotes. '
    'Units in column headings.',
    ['Molecule', 'Method', 'E (Ha)', '\u0394E (kcal/mol)', 'Ref.'],
    [
        ['H\u2082O', 'HF', '-76.0264', '0.00', 'This work'],
        ['H\u2082O', 'CCSD(T)', '-76.3380', '-0.12', 'This work'],
        ['NH\u2083', 'HF', '-56.1956', '0.00', 'This work'],
        ['NH\u2083', 'CCSD(T)', '-56.5630', '-0.08', 'This work'],
        ['CH\u2084', 'HF', '-40.1953', '0.00', 'This work'],
        ['CH\u2084', 'CCSD(T)', '-40.5180', '-0.05', 'This work'],
    ],
    'a) Total electronic energy.\nb) Relative to experimental value.\nc) All calculations performed at equilibrium geometry.'
)

# ══════════════════════════════════════════════════════════════════
# IV. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════
h1('IV. CONCLUSIONS')
body('Summarize the key findings and their implications.')

# ══════════════════════════════════════════════════════════════════
# SUPPLEMENTARY MATERIAL
# ══════════════════════════════════════════════════════════════════
h1('SUPPLEMENTARY MATERIAL')
body(
    'Supplementary material submitted as separate PDF. '
    'S-prefixed numbering (Fig. S1, Table S1, Eq. S1). '
    'All items cited in main text. '
    'After acceptance, deposited in Figshare with DOI.'
)

# ══════════════════════════════════════════════════════════════════
# ACKNOWLEDGMENTS
# ══════════════════════════════════════════════════════════════════
h1('ACKNOWLEDGMENTS')
body('Acknowledge funding, discussions, and contributions. Include grant numbers.')

# ══════════════════════════════════════════════════════════════════
# AUTHOR DECLARATIONS (REQUIRED — 3 parts)
# ══════════════════════════════════════════════════════════════════
h1('AUTHOR DECLARATIONS')
note('[Required: 1) Conflict of Interest (always), 2) Ethics Approval (if animals/humans), '
     '3) Author Contributions (CRediT, always)]')

h2('Conflict of Interest')
body('The authors have no conflicts to disclose.')
note('[COI types — None: "no conflicts to disclose"; '
     'Financial: "[Author] reports grant(s) from {org}..."; '
     'Non-financial: "[Author] is an unpaid member of {org}."; '
     'IP: "[Author] has Patent {number} {pending/issued/licensed}." '
     'Disclose conflicts from 3 years before start of work.]')

h2('Ethics Approval')
note('[Required ONLY if animals/humans. Include committee name, approval ID. '
     'If not applicable, omit or state "Ethics approval not required."]')

h2('Author Contributions')
note('[CRediT taxonomy — REQUIRED. Corresponding author ensures accuracy. '
     'All authors must agree. Published in version of record.]')
p = body('')
r = p.add_run('First Author: '); r.bold = True
p.add_run('Conceptualization (equal); Methodology (lead); Writing – original draft (lead). ')
r = p.add_run('Second Author: '); r.bold = True
p.add_run('Formal analysis (lead); Software (lead); Writing – review & editing (equal). ')
r = p.add_run('Third Author: '); r.bold = True
p.add_run('Supervision (lead); Writing – review & editing (equal).')

# ══════════════════════════════════════════════════════════════════
# DATA AVAILABILITY (REQUIRED)
# ══════════════════════════════════════════════════════════════════
h1('DATA AVAILABILITY')
note('[Select appropriate statement: "authors have data" / "in article" / '
     '"public repo with DOI" / "public repo without DOI" / "no new data" / '
     '"central facility" / "commercial embargo" / "privacy restrictions" / '
     '"third-party restrictions"]')
body(
    'The data that support the findings of this study are available from the corresponding '
    'author upon reasonable request.'
)

# ══════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════
h1('REFERENCES')
note('[JCP: numbered reference style. All cited in text. '
     'No footnotes — use numbered references instead. Include article titles.]')
note('')
for ref in [
    '\u00b9 A. Author, B. Author, and C. Author, "Article title," J. Chem. Phys. 150, 012345 (2019).',
    '',
    '\u00b2 A. B. Author and D. B. Author, "Short title," J. Chem. Phys. 148, 084101 (2018).',
    '',
    '\u00b3 B. Author and D. Author, "Chapter title," in Book Title, edited by E. Editor (Publisher, City, 2020), pp. 100–120.',
    '',
    '\u2074 F. Lastname, "Web resource title," https://example.com (accessed 2024).',
]:
    p = doc.add_paragraph(style='JCP_Ref')
    p.add_run(ref if ref else '')

# ══════════════════════════════════════════════════════════════════
# ALT TEXT PAGE (required upon revision)
# ══════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('ALT TEXT (Required Upon Revision)')
note('[Submit SEPARATE .txt or .docx listing alt text for every figure and table. '
     'If no figures/tables, upload blank file.]')
note('')
note('Alt text principles:')
note('  - Concise: 1-2 sentences, ~25-50 words')
note('  - Specify type (bar graph, scatter plot, flow chart)')
note('  - Describe trends/data, NOT colors/shapes')
note('  - Use discipline-specific terms')
note('  - Do NOT start with "image of" or "figure of"')
note('  - Do NOT repeat caption text')
note('  - Do NOT include opinions')
note('  - Be objective and factual')
note('  - End with a period')
note('')
note('Figure-specific:')
note('  - Graphs: specify type, describe trends, skip colors')
note('  - Diagrams: specify type, describe purpose')
note('  - Complex figures: ONE overall description, no panel-by-panel')
note('  - Decorative elements: NO alt text needed')
note('')
note('Table-specific:')
note('  - Summarize significant trends')
note('  - Identify main point')
note('  - Do NOT write "table of"')
note('  - Do NOT repeat data or headings point-by-point')
note('')

body('Figure 1: [25-50 word description]')
body('Figure 2: [25-50 word description]')
body('Figure 3: [25-50 word description]')
body('Table I: [25-50 word description summarizing significant trends]')
body('Table II: [25-50 word description]')
body('Table III: [25-50 word description]')
body('Table IV: [25-50 word description]')

# ── Page numbers ──
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = run._element
    fldChar1.append(fldChar1.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}))
    run2 = p.add_run()
    instrText = run2._element
    instrText.append(instrText.makeelement(qn('w:instrText'), {}))
    instrText.find(qn('w:instrText')).text = ' PAGE '
    run3 = p.add_run()
    fldChar2 = run3._element
    fldChar2.append(fldChar2.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

output_path = os.path.join(os.path.dirname(__file__), 'JCP_Word_Template.docx')
doc.save(output_path)
print(f'Word template saved to: {output_path}')
